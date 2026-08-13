"""Build editable editorial files required for the RINENG submission."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TITLE = (
    "SABER-PID: Source-Isolated Qualification and Cost-Aware Operation of "
    "Vision-Language Models for P&ID Tag Retrieval"
)
AFFILIATIONS = {
    "a": "Chalmers University of Technology, SE-412 96 Gothenburg, Sweden",
    "b": (
        "Harbin University of Science and Technology, No. 52 Xuefu Road, "
        "Nangang District, Harbin 150080, Heilongjiang, China"
    ),
    "c": "Kiwiar Co., Ltd., 2 Jalan Mat Jambol, #05-30, Singapore 119554, Singapore",
    "d": (
        "Nanjing Agricultural University, No. 1 Weigang, Xuanwu District, "
        "Nanjing 210095, Jiangsu, China"
    ),
    "e": "Hunan University, Lushan South Road, Yuelu District, Changsha 410082, Hunan, China",
}
AUTHORS = (
    {
        "name": "Zhuo Chen",
        "affiliation": "a",
        "email": "zhuoc@chalmers.se",
        "orcid": "0009-0007-7510-8648",
        "designation": "First author; corresponding author",
        "roles": (
            "Conceptualization; Data curation; Formal analysis; Investigation; "
            "Methodology; Project administration; Software; Visualization; "
            "Writing - original draft; Writing - review & editing"
        ),
    },
    {
        "name": "Shuhao Liu",
        "affiliation": "b",
        "email": "2420610137@stu.hrbust.edu.cn",
        "orcid": "0009-0000-2996-9868",
        "designation": "Author",
        "roles": "Data curation; Investigation; Methodology; Validation; Writing - review & editing",
    },
    {
        "name": "Zhi Ling",
        "affiliation": "c",
        "email": "zhi.ling@kiwiar.com",
        "orcid": "0009-0003-2787-156X",
        "designation": "Author",
        "roles": "Investigation; Software; Validation; Writing - review & editing",
    },
    {
        "name": "Yu Yan",
        "affiliation": "c",
        "email": "yu.yan@kiwiar.com",
        "orcid": "0009-0006-4410-4444",
        "designation": "Author",
        "roles": "Data curation; Software; Validation; Writing - review & editing",
    },
    {
        "name": "Qiuxue Wu",
        "affiliation": "b",
        "email": "2420610140@stu.hrbust.edu.cn",
        "orcid": "0009-0001-7290-045X",
        "designation": "Author",
        "roles": "Data curation; Investigation; Validation; Writing - review & editing",
    },
    {
        "name": "Ziyi Kuang",
        "affiliation": "c",
        "email": "ziyi.kuang@kiwiar.com",
        "orcid": "0009-0003-6335-6602",
        "designation": "Author",
        "roles": "Resources; Software; Validation; Writing - review & editing",
    },
    {
        "name": "Zihan Zhao",
        "affiliation": "d",
        "email": "zhaozihan0117@stu.njau.edu.cn",
        "orcid": "0009-0007-7751-7787",
        "designation": "Author",
        "roles": "Formal analysis; Investigation; Validation; Writing - review & editing",
    },
    {
        "name": "Caixin Tan",
        "affiliation": "e",
        "email": "tancaixin2004@hnu.edu.cn",
        "orcid": "0009-0000-8597-3511",
        "designation": "Author",
        "roles": "Investigation; Validation; Visualization; Writing - review & editing",
    },
    {
        "name": "Haiyou Zhang",
        "affiliation": "c",
        "email": "haiyou.zhang@kiwiar.com",
        "orcid": None,
        "designation": "Author",
        "roles": "Project administration; Resources; Supervision; Writing - review & editing",
    },
)
HIGHLIGHTS = (
    "SABER-PID qualifies whether tag output follows the requested P&ID.",
    "Mild image shifts retain 0.548--0.575 requested-drawing F1 effects.",
    "Closest-safe gross-budget InternVL retains a 0.472 drawing effect.",
    "Public DEXPI transfer reaches 0.906 F1; both controls score zero.",
    "Cost ratios select precision-, balance-, or recall-first modes.",
)


def set_font(run, *, size: float, bold: bool = False, color: str = "000000", italic: bool = False) -> None:
    run.font.name = "Calibri"
    r_pr = run._element.get_or_add_rPr()
    r_pr.rFonts.set(qn("w:ascii"), "Calibri")
    r_pr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_exact_spacing(target, *, before: int, after: int, line: int) -> None:
    element = target._p if hasattr(target, "_p") else target._element
    p_pr = element.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    spacing.set(qn("w:line"), str(line))
    spacing.set(qn("w:lineRule"), "auto")


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:keepNext")) is None:
        p_pr.append(OxmlElement("w:keepNext"))


def configure_document(label: str, *, subject: str) -> Document:
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    set_exact_spacing(normal, before=0, after=120, line=264)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 320, 160),
        ("Heading 2", 13, "2E74B5", 240, 120),
        ("Heading 3", 12, "1F4D78", 160, 80),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        set_exact_spacing(style, before=before, after=after, line=264)

    bullet = document.styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    bullet._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    bullet.font.size = Pt(11)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(8)
    bullet.paragraph_format.line_spacing = 1.167

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.add_run(f"RESULTS IN ENGINEERING | {label.upper()}"), size=8.5, bold=True, color="666666")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("SABER-PID submission"), size=8.5, color="777777")

    properties = document.core_properties
    properties.title = f"{label} - SABER-PID"
    properties.subject = subject
    properties.author = "Zhuo Chen"
    properties.keywords = "SABER-PID; Results in Engineering; submission"
    return document


def add_title(document: Document, text: str, subtitle: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(16)
    paragraph.paragraph_format.space_after = Pt(8)
    keep_with_next(paragraph)
    set_font(paragraph.add_run(text), size=23, bold=True)
    if subtitle:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(16)
        keep_with_next(paragraph)
        set_font(paragraph.add_run(subtitle), size=13, color="555555")


def add_label_value(document: Document, label: str, value: str, *, compact: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3 if compact else 6)
    set_font(paragraph.add_run(f"{label}: "), size=10.5, bold=True, color="1F4D78")
    set_font(paragraph.add_run(value), size=10.5)


def build_highlights(output: Path) -> None:
    document = configure_document("Highlights", subject="Article highlights")
    add_title(document, "Highlights", "Research Article | Results in Engineering")
    add_label_value(document, "Manuscript", TITLE)
    heading = document.add_paragraph("Article highlights", style="Heading 1")
    keep_with_next(heading)
    for value in HIGHLIGHTS:
        paragraph = document.add_paragraph(style="List Bullet")
        set_font(paragraph.add_run(value), size=11)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def add_author_line(document: Document, index: int, author: dict[str, str | None]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    set_font(paragraph.add_run(f"{index}. {author['name']}"), size=10.5, bold=True)
    affiliation = paragraph.add_run(str(author["affiliation"]))
    set_font(affiliation, size=8, bold=True, color="1F4D78")
    affiliation.font.superscript = True
    if author["name"] == "Zhuo Chen":
        marker = paragraph.add_run(",*")
        set_font(marker, size=8, bold=True, color="1F4D78")
        marker.font.superscript = True
    set_font(paragraph.add_run(f" - {author['designation']}"), size=10.5, color="555555")
    meta = document.add_paragraph()
    meta.paragraph_format.left_indent = Inches(0.22)
    meta.paragraph_format.space_after = Pt(5)
    orcid = author["orcid"] if author["orcid"] else "Not provided"
    set_font(meta.add_run(f"Email: {author['email']} | ORCID: {orcid}"), size=9.5, color="333333")


def build_title_page(output: Path) -> None:
    document = configure_document("Title Page", subject="Editable title page")
    add_title(document, "Title Page", "Research Article | Results in Engineering")
    heading = document.add_paragraph("Manuscript title", style="Heading 1")
    keep_with_next(heading)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(12)
    set_font(paragraph.add_run(TITLE), size=14, bold=True, color="0B2545")
    heading = document.add_paragraph("Authors", style="Heading 1")
    keep_with_next(heading)
    for index, author in enumerate(AUTHORS, 1):
        add_author_line(document, index, author)
    heading = document.add_paragraph("Affiliations", style="Heading 1")
    keep_with_next(heading)
    for key, value in AFFILIATIONS.items():
        add_label_value(document, key, value, compact=True)
    heading = document.add_paragraph("Corresponding author", style="Heading 1")
    keep_with_next(heading)
    add_label_value(document, "Name", "Zhuo Chen", compact=True)
    add_label_value(document, "Affiliation", AFFILIATIONS["a"], compact=True)
    add_label_value(document, "Email", "zhuoc@chalmers.se", compact=True)
    add_label_value(document, "ORCID", "0009-0007-7510-8648", compact=True)
    heading = document.add_paragraph("Keywords", style="Heading 1")
    keep_with_next(heading)
    document.add_paragraph(
        "P&ID; engineering document intelligence; vision-language model; tag retrieval; "
        "counterfactual evaluation; cost-sensitive decision"
    )
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def build_credit(output: Path) -> None:
    document = configure_document("CRediT Authorship Statement", subject="CRediT author contributions")
    add_title(document, "CRediT Authorship Statement", "Contributor roles for all authors")
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(12)
    set_font(
        paragraph.add_run(
            "The following contributions use the CRediT role names applicable to this work. "
            "Funding acquisition is not assigned because the study received no specific grant."
        ),
        size=10.5,
        color="555555",
    )
    for author in AUTHORS:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.keep_together = True
        set_font(paragraph.add_run(str(author["name"])), size=11, bold=True, color="1F4D78")
        set_font(paragraph.add_run(f" - {author['roles']}."), size=11)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def build_author_information(output: Path) -> None:
    document = configure_document("Author Information", subject="Final author and affiliation metadata")
    add_title(document, "Author Information", "Final author order, affiliations, contacts, and ORCID identifiers")
    add_label_value(document, "Manuscript", TITLE)
    for index, author in enumerate(AUTHORS, 1):
        heading = document.add_paragraph(f"{index}. {author['name']}", style="Heading 2")
        keep_with_next(heading)
        add_label_value(document, "Designation", str(author["designation"]), compact=True)
        add_label_value(
            document,
            "Affiliation",
            f"{author['affiliation']} - {AFFILIATIONS[str(author['affiliation'])]}",
            compact=True,
        )
        add_label_value(document, "Email", str(author["email"]), compact=True)
        add_label_value(document, "ORCID", str(author["orcid"] or "Not provided"), compact=True)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output-dir", default="paper")
    args = parser.parse_args()
    output_dir = Path(args.output_dir)
    outputs = {
        "highlights": output_dir / "Highlights.docx",
        "title_page": output_dir / "Title_Page.docx",
        "credit": output_dir / "CRediT_Authorship_Statement.docx",
        "author_information": output_dir / "Author_Information.docx",
    }
    build_highlights(outputs["highlights"])
    build_title_page(outputs["title_page"])
    build_credit(outputs["credit"])
    build_author_information(outputs["author_information"])
    for name, path in outputs.items():
        print(f"{name}={path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
