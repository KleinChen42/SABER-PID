"""Build the submission-ready Declaration of Interests Word document."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TITLE = (
    "SABER-PID: Source-Isolated Qualification and Cost-Aware Operation of "
    "Vision-Language Models for P&ID Tag Retrieval"
)
AUTHORS = (
    "Zhuo Chen; Shuhao Liu; Zhi Ling; Yu Yan; Qiuxue Wu; Ziyi Kuang; "
    "Zihan Zhao; Caixin Tan; Haiyou Zhang"
)


def set_font(run, *, size: float, bold: bool = False, color: str = "000000") -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table) -> None:
    widths = (2700, 6660)
    table.autofit = False
    table_pr = table._tbl.tblPr
    table_w = table_pr.first_child_found_in("w:tblW")
    table_w.set(qn("w:w"), "9360")
    table_w.set(qn("w:type"), "dxa")
    table_ind = OxmlElement("w:tblInd")
    table_ind.set(qn("w:w"), "120")
    table_ind.set(qn("w:type"), "dxa")
    table_pr.append(table_ind)
    grid = table._tbl.tblGrid
    for old in list(grid):
        grid.remove(old)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_exact_paragraph_spacing(target, *, before: int, after: int, line: int) -> None:
    if hasattr(target, "_p"):
        p_pr = target._p.get_or_add_pPr()
    else:
        p_pr = target._element.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    spacing.set(qn("w:line"), str(line))
    spacing.set(qn("w:lineRule"), "auto")


def build(output: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    set_exact_paragraph_spacing(normal, before=0, after=120, line=264)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_font(header.add_run("RESULTS IN ENGINEERING | DECLARATION"), size=8.5, bold=True, color="666666")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_after = Pt(0)
    set_font(footer.add_run("SABER-PID submission"), size=8.5, color="777777")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(16)
    title.paragraph_format.space_after = Pt(8)
    set_font(title.add_run("Declaration of Interests"), size=23, bold=True, color="000000")

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    set_font(subtitle.add_run("Submission statement for Results in Engineering"), size=13, color="555555")

    table = document.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table)
    rows = (
        ("Manuscript", TITLE),
        ("Authors", AUTHORS),
        ("Corresponding author", "Zhuo Chen (zhuoc@chalmers.se)"),
        ("Prepared", "13 August 2026"),
    )
    for cells, (label, value) in zip(table.rows, rows):
        label_p = cells.cells[0].paragraphs[0]
        value_p = cells.cells[1].paragraphs[0]
        for paragraph in (label_p, value_p):
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.10
        set_font(label_p.add_run(label), size=10.5, bold=True, color="2E4D65")
        set_font(value_p.add_run(value), size=10.5)

    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(18)
    heading.paragraph_format.space_after = Pt(8)
    set_font(heading.add_run("Declaration"), size=16, bold=True, color="2E74B5")
    set_exact_paragraph_spacing(heading, before=320, after=160, line=264)

    statement = document.add_paragraph()
    statement.paragraph_format.space_after = Pt(10)
    set_font(
        statement.add_run(
            "The authors declare that they have no known competing financial "
            "interests or personal relationships that could have appeared to "
            "influence the work reported in this paper."
        ),
        size=11,
    )

    note = document.add_paragraph()
    note.paragraph_format.space_after = Pt(6)
    set_font(
        note.add_run(
            "This declaration is submitted by Zhuo Chen, corresponding author, "
            "on behalf of all listed authors."
        ),
        size=10.5,
        color="555555",
    )

    properties = document.core_properties
    properties.title = "Declaration of Interests — SABER-PID"
    properties.subject = "Competing-interest declaration for Results in Engineering"
    properties.author = "Zhuo Chen"
    properties.keywords = "SABER-PID; declaration of interests; Results in Engineering"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", default="paper/Declaration_of_Interests.docx")
    args = parser.parse_args()
    build(Path(args.output))
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
